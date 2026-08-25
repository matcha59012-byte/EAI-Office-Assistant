"""模块A：企业知识库问答（RAG）——文档管理与单次问答。

会话式问答见 `chat.py`；检索/拒答/LLM 逻辑见 `rag/qa.py`。
"""
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.config import settings
from app.database import get_db
from app.models import Document, KbChunk
from app.rag import qa, splitter, store

router = APIRouter(prefix="/api/kb", tags=["knowledge"])

ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class AskRequest(BaseModel):
    question: str
    library: str | None = None


class SourceItem(BaseModel):
    doc_id: int | None
    chunk_index: int | None
    text: str
    similarity: float


class AskResponse(BaseModel):
    answer: str
    sources: list[SourceItem]
    rejected: bool


class UploadResponse(BaseModel):
    id: int
    title: str
    chunks: int


def _require_admin() -> None:
    if not settings.is_admin:
        raise HTTPException(status_code=403, detail="仅管理员可执行此操作")


def _parse_file(filename: str, content: bytes) -> str:
    """按文件类型抽取纯文本。"""
    name = filename.lower()
    if not any(name.endswith(ext) for ext in ALLOWED_EXTENSIONS):
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {filename}")
    if name.endswith(".txt") or name.endswith(".md"):
        return content.decode("utf-8", errors="ignore")
    if name.endswith(".pdf"):
        from io import BytesIO

        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if name.endswith(".docx"):
        from io import BytesIO

        from docx import Document as DocxDocument

        doc = DocxDocument(BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    raise HTTPException(status_code=400, detail=f"不支持的文件类型: {filename}")


@router.get("/libraries")
def list_libraries(db: Session = Depends(get_db)):
    """资料库列表（文件夹）+ 文档数量。"""
    from sqlalchemy import func

    rows = (
        db.query(Document.library, func.count(Document.id))
        .group_by(Document.library)
        .order_by(Document.library)
        .all()
    )
    return [{"name": name, "count": count} for name, count in rows]


@router.get("/documents")
def list_documents(
    library: str | None = None, db: Session = Depends(get_db)
):
    """只读：员工可浏览文档列表（可按资料库过滤）。"""
    q = db.query(Document)
    if library:
        q = q.filter(Document.library == library)
    rows = q.order_by(Document.created_at.desc()).all()
    return [
        {
            "id": d.id,
            "title": d.title,
            "file_type": d.file_type,
            "library": d.library,
            "created_at": d.created_at,
        }
        for d in rows
    ]


@router.get("/documents/{doc_id}/content")
def document_content(doc_id: int, db: Session = Depends(get_db)):
    """只读：返回文档全文（供中栏查看）。"""
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    if doc.content:
        return {"id": doc.id, "title": doc.title, "content": doc.content}
    # 旧数据兜底：从切片表拼
    chunks = (
        db.query(KbChunk)
        .filter(KbChunk.doc_id == doc_id)
        .order_by(KbChunk.chunk_index.asc())
        .all()
    )
    content = "\n\n".join(c.chunk_text for c in chunks)
    return {"id": doc.id, "title": doc.title, "content": content}


@router.post("/upload", response_model=UploadResponse)
def upload(
    file: UploadFile = File(...),
    library: str = Form("默认资料库"),
    db: Session = Depends(get_db),
):
    """仅管理员可上传。"""
    _require_admin()
    content = file.file.read()
    try:
        text = _parse_file(file.filename, content)
    except HTTPException:
        raise
    except Exception as e:  # noqa: BLE001
        raise HTTPException(status_code=400, detail=f"文件解析失败: {e}")
    if not text.strip():
        raise HTTPException(status_code=400, detail="文件内容为空")

    chunks = splitter.split_text(text)
    if not chunks:
        raise HTTPException(status_code=400, detail="未提取到可检索的文本")

    lib = (library or "默认资料库").strip() or "默认资料库"
    doc = Document(title=file.filename, file_type=file.filename.split(".")[-1], library=lib, content=text)
    db.add(doc)
    db.commit()
    db.refresh(doc)

    store.add_document(doc.id, chunks)
    for i, c in enumerate(chunks):
        db.add(KbChunk(doc_id=doc.id, chunk_index=i, chunk_text=c))
    db.commit()
    return UploadResponse(id=doc.id, title=doc.title, chunks=len(chunks))


@router.delete("/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    """仅管理员可删除。"""
    _require_admin()
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    store.delete_document(doc_id)
    db.query(KbChunk).filter(KbChunk.doc_id == doc_id).delete()
    db.delete(doc)
    db.commit()
    return {"ok": True}


@router.post("/ask", response_model=AskResponse)
def ask(req: AskRequest, db: Session = Depends(get_db)):
    """单次问答（兼容旧接口）；会话式问答见 chat.py。"""
    question = req.question.strip()
    if not question:
        raise HTTPException(status_code=400, detail="问题不能为空")
    # 资料库范围：解析为文档 id 集合
    doc_ids = None
    if req.library:
        doc_ids = [
            d.id for d in db.query(Document).filter(Document.library == req.library).all()
        ]
    try:
        result = qa.answer_question(question, doc_ids=doc_ids)
    except RuntimeError as e:
        raise HTTPException(status_code=502, detail=f"AI 服务调用失败: {e}")
    doc_ids2 = {s["doc_id"] for s in result["sources"] if s.get("doc_id")}
    titles: dict = {}
    if doc_ids2:
        rows = db.query(Document).filter(Document.id.in_(doc_ids2)).all()
        titles = {d.id: d.title for d in rows}
    answer = qa.append_sources_footer(result["answer"], result["sources"], titles)
    return AskResponse(
        answer=answer,
        sources=[SourceItem(**s) for s in result["sources"]],
        rejected=result["rejected"],
    )
